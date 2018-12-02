import os
import re

from docx import Document

from metacorps.projects.common.export_project import ProjectExporter
from iatv.iatv import Show


def download_instance_transcripts(debug_lim=3,
                                  write_dir='transcripts'):

    df = ProjectExporter('EPA Metvi').export_dataframe()

    inst_ids = df.iatv_id.unique()
    if debug_lim > 0:
        inst_ids = inst_ids[:debug_lim]

    N_inst = len(inst_ids)

    for idx, inst_id in enumerate(inst_ids):

        try:
            show = Show(inst_id)
            trans = show.get_transcript(verbose=False)

            fulltext = str(u'\n\n'.join(trans).encode('utf-8'))

            write_path = os.path.join(write_dir, inst_id + '.txt')
            open(write_path, 'w').write(fulltext)

            print('saved {} to {} ({}/{})'.format(inst_id, write_path, idx+1, N_inst))

        except:
            print('failed to save {} to {} ({}/{})'.format(inst_id, write_path, idx+1, N_inst))


def format_snippet(transcript, re_word=r'STRANGL'):

    f_iter = re.finditer(re_word, transcript)
    tr = transcript

    def cleanup(s):
        return s.replace("\\'", "'").replace('\\n\\n', '\n\n')

    formatted = []
    for idx, m in enumerate(f_iter):
        # pre = '# {}\n\n'.format(idx + 1) + cleanup(tr[m.start() - 1000: m.start() - 65]).lower()
        pre = cleanup(tr[m.start() - 1000: m.start() - 65]).lower()
        focus = cleanup(tr[m.start() - 65: m.end() + 65])
        post = cleanup(tr[m.end() + 65: m.end() + 1000]).lower()
        # formatted.append(pre + focus + post)

    # return formatted
    return (pre, focus, post)


def make_docx(transcript_paths,
              title='EPA Metvi snippets',
              docx_path='transcripts.docx'):

    docx = Document()
    docx.add_heading(title, 0)

    transcript_paths.sort(key=lambda x: int(x.split('_')[1]))

    for trp in transcript_paths:

        tr = open(trp, 'r').read()
        pre, focus, post = format_snippet(tr)

        split_path = trp.split('_')
        channel = split_path[0].split('/')[-1]
        year = split_path[1][:4]
        month = split_path[1][4:6]
        day = split_path[1][6:]

        show = ' '.join(split_path[3:]).replace('.txt', '')

        docx.add_heading('{} - {} - {}/{}/{}'.format(
            channel, show, month, day, year
        ))

        docx.add_paragraph(
            'https://archive.org/details/' +
            trp.split('/')[-1].replace('.txt', '')
        )

        p = docx.add_paragraph(pre)
        p.add_run(focus).bold = True
        p.add_run(post)

    docx.save(docx_path)
